import subprocess
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import ttkbootstrap as ttk
import time
import threading
from docx import Document
import os
from PIL import Image, ImageTk
import sys
import webbrowser

class Task:
    def __init__(self, parent, tasks, task_no, task_name=None, time=15, test_time=5, break_time=10):
        self.parent = parent
        self.tasks = tasks
        self.task_no = task_no
        self.task_name = task_name
        self.main_time = time
        self.test_time = test_time
        self.break_time = break_time

        self.time_remaining = self.main_time * 60
        self.timer_running = False
        self.current_timer = "main"  # Can be 'main', 'test', 'break'
        self.pdf_path = ''

        self.create_widgets()

    def create_widgets(self):
        # frame
        self.frame = ttk.Frame(self.parent,
                               padding=10,
                               relief="solid")
        self.frame.pack(fill="x", padx=5, pady=5)

        self.task_no_label = ttk.Label(self.frame,
                                       text=self.task_no,
                                       font=("Helvetica", 16))
        self.task_no_label.grid(row=0, column=0, padx=5, sticky="w")

        self.main_timer_label = ttk.Label(self.frame, text="Main Time (mins):", font=("Helvetica", 12))
        self.main_timer_label.grid(row=0, column=1, padx=5, sticky="e")
        self.timer_label = ttk.Label(self.frame,
                                     text="15:00",
                                     font=("Helvetica", 16),
                                     bootstyle="info")
        self.timer_label.grid(row=0, column=2)

        self.task_name_label = ttk.Label(self.frame,
                                         text="Task Name",
                                         font=("Helvetica", 12))
        self.task_name_label.grid(row=0, column=3, padx=5, sticky="w")

        self.task_entry = ttk.Entry(self.frame,
                                    font=("Helvetica", 14),
                                    width=10,
                                    bootstyle='light')
        self.task_entry.grid(row=0, column=4, padx=6)

        self.edit_button = ttk.Button(self.frame,
                                       text="Edit Timers",
                                       command=self.edit_timers,
                                       bootstyle='info',
                                       cursor='hand2')
        self.edit_button.grid(row=0, column=5, padx=5)

        self.start_button = ttk.Button(self.frame,
                                       text="Start",
                                       command=self.start_timer,
                                       bootstyle="success",
                                       cursor='hand2')
        self.start_button.grid(row=1, column=0, padx=10, pady=10)

        self.reset_button = ttk.Button(self.frame,
                                       text="Reset",
                                       command=self.reset_timer,
                                       bootstyle='info',
                                       cursor='hand2')
        self.reset_button.grid(row=1, column=1, padx=10, pady=10)

        self.delete_button = ttk.Button(self.frame,
                                        text="Delete",
                                        command=self.delete_timer,
                                        bootstyle="danger",
                                        cursor='hand2')
        self.delete_button.grid(row=1, column=2, padx=10, pady=10)

        self.open_pdf_button = ttk.Button(self.frame,
                                          text="Open PDF",
                                          command=self.open_pdf,
                                          bootstyle="info",
                                          cursor='hand2')
        self.open_pdf_button.grid(row=1, column=3, padx=5, pady=5)

        self.pdf_label = ttk.Label(self.frame,
                                   text="No PDF selected",
                                   font=("Helvetica", 12))
        self.pdf_label.grid(row=1, column=4, padx=5, pady=5)

    def start_timer(self):
        if not self.timer_running:
            self.timer_running = True
            threading.Thread(target=self.run_timer, daemon=True).start()

    def reset_timer(self):
        self.timer_running = False
        self.current_timer = "main"
        self.time_remaining = self.main_time * 60
        self.update_timer_label()

    def delete_timer(self):
        self.frame.destroy()
        self.tasks.remove(self)

    def run_timer(self):
        while self.time_remaining > 0 and self.timer_running:
            mins, secs = divmod(self.time_remaining, 60)
            self.timer_label.config(text=f"{mins:02}:{secs:02}")
            time.sleep(1)
            self.time_remaining -= 1

        if self.time_remaining == 0:
            if self.current_timer == "main":
                self.current_timer = "test"
                self.time_remaining = self.test_time * 60
                messagebox.showinfo("Main Time's Up!", "Starting Test Timer.")
                self.open_word_file()
                self.main_timer_label.config(text="Test Time (mins):")
                self.run_timer()
            elif self.current_timer == "test":
                self.current_timer = "break"
                self.time_remaining = self.break_time * 60
                messagebox.showinfo("Test Time's Up!", "Starting Break Timer.")
                self.main_timer_label.config(text="Break Time (mins):")
                self.run_timer()
            elif self.current_timer == "break":
                self.timer_running = False
                messagebox.showinfo("Break Time's Up!", "All timers completed!")
                self.frame.destroy()


    def open_word_file(self):
        if self.task_entry.get() != '':
            self.task_name = self.task_entry.get()
        else:
            self.task_name = self.pdf_path
        doc = Document()
        doc.add_heading(f"{self.task_no}")

        doc.add_paragraph(f"Task Name: {self.task_name}")

        doc_file = f"{self.task_name.replace(" ", '')}.docx"
        doc.save(doc_file)
        try:
            subprocess.run(["start", doc_file], shell=True, check=True)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to open Word file: {e}")
    def update_timer_label(self):
        mins, secs = divmod(self.time_remaining, 60)
        self.timer_label.config(text=f"{mins:02}:{secs:02}")

    def open_pdf(self):
        self.pdf_path = filedialog.askopenfilename(filetypes=[("PDF Files", "*.pdf")])
        if self.pdf_path:
            try:
                webbrowser.open(self.pdf_path)
                self.pdf_name = os.path.basename(self.pdf_path)
                max_length = 20
                if len(self.pdf_name) > max_length:
                    self.pdf_name = self.pdf_name[:max_length] + "..."
                self.pdf_label.config(text=self.pdf_name)
            except Exception as e:
                messagebox.showerror("Error", f"Failed to open PDF: {e}")

    def edit_timers(self):
        def save_changes():
            try:
                self.main_time = int(main_time_var.get())
                self.test_time = int(test_time_var.get())
                self.break_time = int(break_time_var.get())
                self.reset_timer()
                edit_window.destroy()
            except ValueError:
                messagebox.showerror("Error", "Please enter valid numbers!")

        edit_window = tk.Toplevel(self.parent)
        edit_window.title("Edit Timers")

        ttk.Label(edit_window, text="Main Time (mins):").grid(row=0, column=0, padx=5, pady=5)
        main_time_var = tk.StringVar(value=str(self.main_time))
        ttk.Entry(edit_window, textvariable=main_time_var).grid(row=0, column=1, padx=5, pady=5)

        ttk.Label(edit_window, text="Test Time (mins):").grid(row=1, column=0, padx=5, pady=5)
        test_time_var = tk.StringVar(value=str(self.test_time))
        ttk.Entry(edit_window, textvariable=test_time_var).grid(row=1, column=1, padx=5, pady=5)

        ttk.Label(edit_window, text="Break Time (mins):").grid(row=2, column=0, padx=5, pady=5)
        break_time_var = tk.StringVar(value=str(self.break_time))
        ttk.Entry(edit_window, textvariable=break_time_var).grid(row=2, column=1, padx=5, pady=5)

        ttk.Button(edit_window, text="Save", command=save_changes).grid(row=3, column=0, columnspan=2, pady=10)

class PomodoroApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Pomodoro Timer")
        self.root.geometry("1280x720")
        self.root.resizable(False, False)

        self.header = ttk.Label(self.root, text="Pomodoro Timer", font=("Helvetica", 24, "bold underline"), bootstyle="Danger")
        self.header.pack(pady=20)

        self.root.img = tk.PhotoImage(file=resource_path("logo.png"))

        self.logo = ttk.Label(self.root, image=self.root.img)
        self.logo.pack(pady=20)
        self.tagline = ttk.Label(self.root, text="Time in your hands", font=("Helvetica", 14, "italic bold "), bootstyle="Danger")
        self.tagline.pack()

        self.tutorial = ttk.Label(self.root, text="Click here to get started", font=("Helvetica", 14, "italic bold "), bootstyle="Danger")
        self.tutorial.pack(pady=20)

        arrow_image = Image.open(resource_path("arrow.png"))
        arrow_image_resized = arrow_image.resize((50, 100))  # Resize to 50x100
        self.root.arrow_image_tk = ImageTk.PhotoImage(arrow_image_resized)
        self.arrow_label = ttk.Label(self.root, image=self.root.arrow_image_tk)
        self.arrow_label.pack()

        # Task container
        self.task_container = ttk.Frame(self.root)
        self.task_container.pack(fill="both",
                                 expand=True,
                                 padx=10,
                                 pady=10)

        # Add Task Button
        self.add_task_button = ttk.Button(self.root,
                                          width=20,
                                          text="Add New Task",
                                          command=self.add_task,
                                          bootstyle="info",
                                          cursor='hand2',
                                          )
        self.add_task_button.pack(side="bottom",
                                  padx=10,
                                  pady=10)

        self.temp = 1
        self.tasks = []

    def add_task(self):
        task_no = f"Task {self.temp}"
        if self.temp == 1:
            self.header.forget()
            self.logo.forget()
            self.tagline.forget()
            self.tutorial.forget()
            self.arrow_label.forget()
        self.temp += 1
        new_task = Task(self.task_container,
                        self.tasks,
                        task_no)
        self.tasks.append(new_task)

def resource_path(relative_path):
    try:
        base_path = getattr(sys, '_MEIPASS', os.path.abspath('.'))
    except Exception:
        base_path = os.path.abspath('.')
    return os.path.join(base_path, relative_path)

if __name__ == "__main__":
    root = ttk.Window("Pomodoro Timer", "superhero")
    app = PomodoroApp(root)
    root.mainloop()
